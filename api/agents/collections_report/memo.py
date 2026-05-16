"""Performance memo generator.

Produces a structured Memo object that can be rendered to:
- plain text (for the in-app modal preview)
- PDF (reportlab)
- .docx (python-docx — opens cleanly in Google Docs and Microsoft Word)

The structure is the single source of truth; renderers below transform it.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from decimal import Decimal
from io import BytesIO
from typing import Literal, Optional

from api.agents.collections_report.engine import ReportData


def _ghs(n) -> str:
    return f"GHS {float(n):,.2f}"


def _pct(n) -> str:
    return f"{float(n) * 100:.1f}%"


# ---------------------------------------------------------------------------
# Memo data model
# ---------------------------------------------------------------------------


@dataclass
class MemoBullet:
    text: str
    emphasis: bool = False  # bold/bigger for headline items


@dataclass
class MemoRiderRow:
    rank: int
    name: str
    customer_id: str
    outstanding_ghs: float
    open_invoices: int
    band: str
    months_since: int


@dataclass
class MemoSection:
    title: str
    kind: Literal["bullets", "riders", "actions", "notes"]
    bullets: list[MemoBullet] = field(default_factory=list)
    riders: list[MemoRiderRow] = field(default_factory=list)
    actions: list[str] = field(default_factory=list)
    notes: list[str] = field(default_factory=list)
    subheading: Optional[str] = None


@dataclass
class Memo:
    title: str
    period_label: str
    scope_label: str
    prepared_on: str
    sections: list[MemoSection] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Builder (turns ReportData → Memo)
# ---------------------------------------------------------------------------


def build_memo(report: ReportData) -> Memo:
    r = report
    status_label = {
        "active": "Active book",
        "recovery": "Recovery (churned)",
        "completed": "Completed riders",
        "all": "All riders",
    }.get(r.status_filter, r.status_filter)

    memo = Memo(
        title="Wahu Collections Performance Memo",
        period_label=r.window_label,
        scope_label=f"{status_label}  ·  Fleet: {r.fleet_filter}",
        prepared_on=r.as_of.isoformat(),
    )

    # ---- Portfolio status ----
    ratio = (
        r.lifetime_collected_ghs / r.lifetime_invoiced_ghs
        if r.lifetime_invoiced_ghs > 0
        else Decimal("0")
    )
    status = MemoSection(title="Portfolio status", kind="bullets")
    status.bullets += [
        MemoBullet(f"Riders in scope: {r.active_riders} of {r.total_rider_population} total", emphasis=True),
        MemoBullet(f"Invoiced (in scope): {_ghs(r.lifetime_invoiced_ghs)}", emphasis=True),
        MemoBullet(f"Collected (in scope): {_ghs(r.lifetime_collected_ghs)} ({_pct(ratio)})", emphasis=True),
        MemoBullet(f"Outstanding: {_ghs(r.lifetime_outstanding_ghs)}", emphasis=True),
        MemoBullet(f"Open invoice lines: {r.open_invoice_lines}"),
    ]
    if r.cash_in_window_ghs:
        status.bullets.append(
            MemoBullet(f"Cash received this window: {_ghs(r.cash_in_window_ghs)}")
        )
        if r.cash_applied_to_prior_ghs:
            status.bullets.append(
                MemoBullet(
                    f"   of which paying down older invoices: {_ghs(r.cash_applied_to_prior_ghs)}"
                )
            )
    memo.sections.append(status)

    # ---- Risk signal ----
    band_lookup = {b.band: b for b in r.risk_bands}
    risk_d = band_lookup.get("D")
    risk_e = band_lookup.get("E")
    band_at_risk = (risk_d.riders if risk_d else 0) + (risk_e.riders if risk_e else 0)
    outstanding_at_risk = (
        (risk_d.outstanding_ghs if risk_d else Decimal("0"))
        + (risk_e.outstanding_ghs if risk_e else Decimal("0"))
    )

    risk = MemoSection(title="Risk signal", kind="bullets")
    if band_at_risk:
        risk.bullets.append(
            MemoBullet(
                f"{band_at_risk} riders in bands D/E carrying {_ghs(outstanding_at_risk)} outstanding — these are the active ones drifting toward churn.",
                emphasis=True,
            )
        )
    else:
        risk.bullets.append(MemoBullet("No riders in bands D/E. Portfolio is healthy on the lifetime ratio."))

    aged = [a for a in r.ageing if a.label not in ("Current (0–30d)",)]
    aged_total = sum((a.outstanding_ghs for a in aged), Decimal("0"))
    aged_count = sum(a.open_invoices for a in aged)
    if aged_total > 0:
        risk.bullets.append(
            MemoBullet(
                f"{aged_count} open invoices over 30 days worth {_ghs(aged_total)} — needs SMS dunning or field follow-up."
            )
        )
    over_90 = [a for a in r.ageing if a.label in ("91–180d", "181–365d", "365d+")]
    over_90_total = sum((a.outstanding_ghs for a in over_90), Decimal("0"))
    over_90_count = sum(a.open_invoices for a in over_90)
    if over_90_count:
        risk.bullets.append(
            MemoBullet(
                f"{over_90_count} invoices over 90 days worth {_ghs(over_90_total)} — escalate to recovery.",
                emphasis=True,
            )
        )
    memo.sections.append(risk)

    # ---- Top problem accounts ----
    top_outstanding = sorted(
        [s for s in r.scorecards if s.lifetime_outstanding_ghs > 0],
        key=lambda s: s.lifetime_outstanding_ghs,
        reverse=True,
    )[:10]
    if top_outstanding:
        combined = sum(
            (s.lifetime_outstanding_ghs for s in top_outstanding), Decimal("0")
        )
        section = MemoSection(
            title=f"Top {len(top_outstanding)} by outstanding",
            subheading=f"{_ghs(combined)} combined",
            kind="riders",
        )
        for i, s in enumerate(top_outstanding, start=1):
            section.riders.append(
                MemoRiderRow(
                    rank=i,
                    name=s.customer_name,
                    customer_id=s.customer_id,
                    outstanding_ghs=float(s.lifetime_outstanding_ghs),
                    open_invoices=s.open_invoices,
                    band=s.risk_band,
                    months_since=s.months_since_last_invoice,
                )
            )
        memo.sections.append(section)

    # ---- Priority actions ----
    actions: list[str] = []
    if top_outstanding:
        actions.append(
            f"Call the {len(top_outstanding)} top-outstanding accounts above — combined {_ghs(combined)}."
        )
    if over_90_count:
        actions.append(
            f"Field visit for invoices over 90 days ({over_90_count} invoices, {_ghs(over_90_total)})."
        )
    if risk_e and risk_e.riders:
        actions.append(
            f"Band E intervention plan — {risk_e.riders} riders with <30% paid lifetime."
        )
    aged_31_60 = next((a for a in r.ageing if a.label == "31–60d"), None)
    if aged_31_60 and aged_31_60.open_invoices:
        actions.append(
            f"SMS dunning for 31–60d bucket ({aged_31_60.open_invoices} invoices, {_ghs(aged_31_60.outstanding_ghs)}) — next push: Wed/Fri."
        )
    if r.status_filter == "recovery" and r.active_riders:
        actions.append(
            f"Recovery outreach for {r.active_riders} churned riders with outstanding — propose settlement plans."
        )
    if r.status_filter == "completed" and r.active_riders:
        actions.append(
            f"Issue completion certificates for {r.active_riders} fully-paid-out riders."
        )
    if not actions:
        actions.append("No urgent actions — portfolio is current. Maintain weekly cadence.")
    memo.sections.append(MemoSection(title="Priority actions for this period", kind="actions", actions=actions))

    # ---- Notes ----
    notes_section = MemoSection(title="Notes", kind="notes")
    notes_section.notes.append(
        f"The collection ratio above ({_pct(ratio)}) reflects cash applied to invoices issued in scope. "
        "Cash that came in this period but paid down older invoices is captured under 'Cash received this window'. "
        "Older invoice exports in Drive may need refreshing to capture payments made this period against earlier invoices."
    )
    notes_section.notes.append("See the Portfolio Trends dashboard for month-over-month context and rider rankings.")
    memo.sections.append(notes_section)

    return memo


# ---------------------------------------------------------------------------
# Renderers
# ---------------------------------------------------------------------------


def render_text(memo: Memo) -> str:
    """Plain text — what the in-app modal displays. No markdown syntax."""
    lines: list[str] = []
    lines.append(memo.title)
    lines.append(f"Period: {memo.period_label}")
    lines.append(f"Scope:  {memo.scope_label}")
    lines.append(f"Prepared: {memo.prepared_on}")
    lines.append("")
    for s in memo.sections:
        title = s.title
        if s.subheading:
            title += f"  ({s.subheading})"
        lines.append(title)
        lines.append("-" * min(60, max(20, len(title))))
        if s.kind == "bullets":
            for b in s.bullets:
                lines.append(f"  • {b.text}")
        elif s.kind == "riders":
            for r in s.riders:
                lines.append(
                    f"  {r.rank:>2}. {r.name} ({r.customer_id}) — "
                    f"GHS {r.outstanding_ghs:,.2f} outstanding · "
                    f"{r.open_invoices} open · band {r.band} · "
                    f"{r.months_since}mo since last invoice"
                )
        elif s.kind == "actions":
            for i, a in enumerate(s.actions, start=1):
                lines.append(f"  {i}. {a}")
        elif s.kind == "notes":
            for n in s.notes:
                lines.append(f"  • {n}")
        lines.append("")
    return "\n".join(lines)


def render_pdf(memo: Memo) -> bytes:
    """A4 PDF via reportlab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether,
    )

    INK = colors.HexColor("#0F1419")
    MUTED = colors.HexColor("#4A5159")
    ACCENT = colors.HexColor("#8B5A2B")
    LINE = colors.HexColor("#E8E2D6")

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=memo.title,
    )

    base = getSampleStyleSheet()
    H1 = ParagraphStyle("H1", parent=base["Heading1"], fontSize=18, leading=22,
                        textColor=INK, spaceAfter=4, fontName="Helvetica-Bold")
    H2 = ParagraphStyle("H2", parent=base["Heading2"], fontSize=12, leading=16,
                        textColor=INK, spaceBefore=10, spaceAfter=4, fontName="Helvetica-Bold")
    sub = ParagraphStyle("sub", parent=base["Normal"], fontSize=10, leading=14,
                         textColor=MUTED, italic=True, spaceAfter=10)
    body = ParagraphStyle("body", parent=base["Normal"], fontSize=10.5, leading=15,
                          textColor=INK, fontName="Helvetica")
    bold = ParagraphStyle("bold", parent=body, fontName="Helvetica-Bold")
    bullet_style = ParagraphStyle("bullet", parent=body, leftIndent=14, bulletIndent=2)

    story: list = []
    story.append(Paragraph(memo.title, H1))
    story.append(Paragraph(
        f"<b>Period:</b> {memo.period_label}<br/>"
        f"<b>Scope:</b> {memo.scope_label}<br/>"
        f"<b>Prepared:</b> {memo.prepared_on}",
        sub,
    ))

    for s in memo.sections:
        title = s.title + (f"  ({s.subheading})" if s.subheading else "")
        story.append(Paragraph(title, H2))

        if s.kind == "bullets":
            for b in s.bullets:
                style = bold if b.emphasis else bullet_style
                story.append(Paragraph(f"• {b.text}", style))
        elif s.kind == "actions":
            for i, a in enumerate(s.actions, start=1):
                story.append(Paragraph(f"{i}. {a}", body))
        elif s.kind == "notes":
            for n in s.notes:
                story.append(Paragraph(n, body))
                story.append(Spacer(1, 3))
        elif s.kind == "riders":
            data = [["#", "Rider", "Outstanding", "Open", "Band", "Mo since"]]
            for r in s.riders:
                data.append([
                    str(r.rank),
                    f"{r.name}\n{r.customer_id}",
                    f"GHS {r.outstanding_ghs:,.2f}",
                    str(r.open_invoices),
                    r.band,
                    f"{r.months_since}mo",
                ])
            tbl = Table(data, colWidths=[10 * mm, 70 * mm, 35 * mm, 15 * mm, 15 * mm, 20 * mm])
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), INK),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ALIGN", (0, 0), (0, -1), "RIGHT"),
                ("ALIGN", (2, 0), (-1, -1), "RIGHT"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LINEBELOW", (0, 0), (-1, -1), 0.25, LINE),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#FAF7F2")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]))
            story.append(tbl)

    doc.build(story)
    return buf.getvalue()


def render_docx(memo: Memo) -> bytes:
    """Word .docx — opens cleanly in Google Docs (File → Open → Upload)."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm

    INK = RGBColor(0x0F, 0x14, 0x19)
    MUTED = RGBColor(0x4A, 0x51, 0x59)

    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Title
    title = doc.add_heading(memo.title, level=0)
    for run in title.runs:
        run.font.color.rgb = INK

    p = doc.add_paragraph()
    for label, value in (
        ("Period: ", memo.period_label),
        ("\nScope:  ", memo.scope_label),
        ("\nPrepared: ", memo.prepared_on),
    ):
        run_lbl = p.add_run(label)
        run_lbl.bold = True
        run_lbl.font.size = Pt(10)
        run_lbl.font.color.rgb = MUTED
        run_val = p.add_run(value)
        run_val.font.size = Pt(10)
        run_val.font.color.rgb = MUTED

    for s in memo.sections:
        h = doc.add_heading(
            s.title + (f"  ({s.subheading})" if s.subheading else ""),
            level=2,
        )
        for run in h.runs:
            run.font.color.rgb = INK

        if s.kind == "bullets":
            for b in s.bullets:
                para = doc.add_paragraph(b.text, style="List Bullet")
                if b.emphasis:
                    for run in para.runs:
                        run.bold = True
        elif s.kind == "actions":
            for a in s.actions:
                doc.add_paragraph(a, style="List Number")
        elif s.kind == "notes":
            for n in s.notes:
                doc.add_paragraph(n)
        elif s.kind == "riders":
            tbl = doc.add_table(rows=1, cols=6)
            tbl.style = "Light Grid Accent 1"
            hdr_cells = tbl.rows[0].cells
            for i, h in enumerate(["#", "Rider", "Outstanding", "Open", "Band", "Mo since"]):
                hdr_cells[i].text = h
                for run in hdr_cells[i].paragraphs[0].runs:
                    run.bold = True
            for r in s.riders:
                row = tbl.add_row().cells
                row[0].text = str(r.rank)
                row[1].text = f"{r.name}\n{r.customer_id}"
                row[2].text = f"GHS {r.outstanding_ghs:,.2f}"
                row[3].text = str(r.open_invoices)
                row[4].text = r.band
                row[5].text = f"{r.months_since}mo"

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
